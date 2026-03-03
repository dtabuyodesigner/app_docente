import os
import re
from docx import Document
from docx.shared import Inches

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} no existe.")
        return

    doc = Document()
    doc.add_heading('Walkthrough: APP_EVALUAR - Blindaje Técnico y Mejoras de Biblioteca', 0)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        
        # Images (special handling for absolute paths)
        elif line.startswith('![') and ']' in line:
            img_match = re.search(r'\((.*?)\)', line)
            if img_match:
                img_path = img_match.group(1)
                if os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(6))
                else:
                    doc.add_paragraph(f"[Imagen no encontrada: {img_path}]")
        
        # List items
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            # Simple numeric list handling
            doc.add_paragraph(line[3:], style='List Number')
        
        # Plain text
        else:
            # Handle minor formatting like **bold** (very basic)
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)

    doc.save(docx_path)
    print(f"Documento guardado en: {docx_path}")

if __name__ == "__main__":
    md_file = "/home/danito73/.gemini/antigravity/brain/7fde250b-fd88-4b53-90bc-6e4887853e20/walkthrough.md"
    docx_file = "/home/danito73/Documentos/APP_EVALUAR/Walkthrough_Navegacion_Material_v3.docx"
    convert_md_to_docx(md_file, docx_file)
